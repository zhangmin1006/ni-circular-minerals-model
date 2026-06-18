"""
Export the model's Markdown outputs (the Q2.1-2.7 findings memos) and the
top-level briefing docs to Word (.docx), into ../word/.

A lightweight Markdown -> docx converter (python-docx) covering exactly the
features these files use: headings, tables, bullet/numbered lists, blockquotes,
code fences, **bold**, `inline code`, and [text](link). No external tools needed.

Run:  python export_word.py
"""

import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(__file__)
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT_DIR = os.path.join(ROOT, "word")

# (source markdown, output .docx name)
FILES = [
    (os.path.join(ROOT, "NI_MINERALS_MODEL_REPORT.md"), "NI_Minerals_Model_Report.docx"),
    (os.path.join(ROOT, "NI_MINERALS_TECHNICAL_REPORT.md"), "NI_Minerals_Technical_Report.docx"),
    (os.path.join(HERE, "outputs", "q2_1_memo.md"), "Q2.1_circularity_innovation.docx"),
    (os.path.join(HERE, "outputs", "q2_2_memo.md"), "Q2.2_opportunities_challenges.docx"),
    (os.path.join(HERE, "outputs", "q2_3_memo.md"), "Q2.3_business_support.docx"),
    (os.path.join(HERE, "outputs", "q2_4_memo.md"), "Q2.4_secure_supply.docx"),
    (os.path.join(HERE, "outputs", "q2_5_memo.md"), "Q2.5_employment_skills.docx"),
    (os.path.join(HERE, "outputs", "q2_6_memo.md"), "Q2.6_economic_benefits.docx"),
    (os.path.join(HERE, "outputs", "q2_7_memo.md"), "Q2.7_negative_impacts.docx"),
    (os.path.join(HERE, "outputs", "q_demand_supply_memo.md"), "Demand_supply_analysis.docx"),
    (os.path.join(ROOT, "EXECUTIVE_SUMMARY.md"), "Executive_Summary.docx"),
    (os.path.join(ROOT, "TECHNICAL_DOCUMENTATION.md"), "Technical_Documentation.docx"),
]

_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_IMAGE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def _add_runs(paragraph, text):
    """Render inline markdown (**bold**, `code`, [text](url)) into runs."""
    text = _LINK.sub(lambda m: m.group(1) + (f" ({m.group(2)})"
                     if m.group(2).startswith("http") else ""), text)
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        else:
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def _is_separator(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def _cells(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def md_to_docx(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # code fences ``` ... ```
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                code_buf = []
            in_code = not in_code
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # tables
        if _is_table_row(line):
            block = []
            while i < len(lines) and _is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            rows = [_cells(b) for b in block if not _is_separator(b)]
            if rows:
                ncol = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=ncol)
                table.style = "Light Grid Accent 1"
                table.autofit = True
                for ridx, row in enumerate(rows):
                    cells = table.add_row().cells
                    for c in range(ncol):
                        txt = row[c] if c < len(row) else ""
                        para = cells[c].paragraphs[0]
                        _add_runs(para, txt)
                        for run in para.runs:
                            run.font.size = Pt(8)
                            if ridx == 0:
                                run.bold = True
            continue

        # images:  ![caption](relative/path.png)
        mi = _IMAGE.match(stripped)
        if mi:
            caption, rel = mi.group(1), mi.group(2)
            img = rel if os.path.isabs(rel) else os.path.join(os.path.dirname(md_path), rel)
            if os.path.exists(img):
                doc.add_picture(img, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(caption)
                    r.italic = True
                    r.font.size = Pt(8.5)
            else:
                doc.add_paragraph(f"[missing figure: {rel}]")
            i += 1
            continue

        # headings
        m = re.match(r"(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 4))
            _add_runs(h, m.group(2))
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|\*{3,}", stripped):
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # numbered list
        if re.match(r"\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # bullet list
        if re.match(r"[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        _add_runs(p, stripped)
        i += 1

    doc.save(docx_path)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    done, locked = 0, []
    for md_path, name in FILES:
        if not os.path.exists(md_path):
            print(f"  skip (missing): {md_path}")
            continue
        out = os.path.join(OUT_DIR, name)
        try:
            md_to_docx(md_path, out)
        except PermissionError:
            # the .docx is open (Word/OneDrive lock) — skip it, keep going so one
            # open document cannot block the rest of the export.
            locked.append(name)
            print(f"  SKIP (file locked/open, close it and re-run): {name}")
            continue
        print(f"  wrote {os.path.relpath(out, ROOT)}")
        done += 1
    print(f"\n{done} Word documents written to {os.path.relpath(OUT_DIR, ROOT)}/")
    if locked:
        print(f"{len(locked)} skipped because open elsewhere: {', '.join(locked)}")


if __name__ == "__main__":
    main()
