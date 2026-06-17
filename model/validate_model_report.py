"""
Run model validation and write a Word report.

This script wraps the existing verification harness in ``verify_model.py`` and
turns the results into durable validation artefacts:

  - outputs/model_validation_results.csv
  - outputs/model_validation_results.json
  - ../word/Model_Validation_Report.docx

Run from the repository root or from ``model/``:

    python model/validate_model_report.py
"""

import contextlib
import io
import json
import os
import sys
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
OUT = os.path.join(HERE, "outputs")
WORD = os.path.join(ROOT, "word")
SRC = os.path.join(HERE, "src")

if HERE not in sys.path:
    sys.path.insert(0, HERE)
if SRC not in sys.path:
    sys.path.insert(0, SRC)

import verify_model  # noqa: E402


TEST_GROUPS = [
    ("Calibration to external anchors", verify_model.test_validation),
    ("Material-flow accounting and supply closure", verify_model.test_mass_balance),
    ("Supply-share bounds", verify_model.test_share_closure),
    ("Output integrity", verify_model.test_no_nan_negative),
    ("Determinism and repeatability", verify_model.test_determinism),
    ("SAM and CGE benchmark replication", verify_model.test_sam_cge),
    ("Spatial allocation", verify_model.test_spatial),
    ("Strategic stockpile behaviour", verify_model.test_stockpile),
    ("Company-register integrity", verify_model.test_company_register),
    ("Economic sanity ranges", verify_model.test_economic_sanity),
    ("Geopolitical-shock behaviour", verify_model.test_geopolitical),
    ("Employment and skills layer", verify_model.test_employment),
    ("Negative-impact layer", verify_model.test_impact),
    ("Property-based stress tests", verify_model.test_fuzz),
]


VALIDATION_DESIGN = [
    {
        "layer": "External calibration",
        "purpose": "Check that the Input-Output core reproduces Minviro scenario anchors.",
        "evidence": "One-mine and two-mine output, jobs and direct mining GVA are compared against published anchors.",
        "pass_rule": "Output within GBP0.15m, jobs within 1 job, direct mining GVA within 5%.",
    },
    {
        "layer": "Physical consistency",
        "purpose": "Confirm that the Material Flow Account conserves material under baseline and disrupted supply.",
        "evidence": "Domestic supply, imports, secondary recovery and unmet demand are checked for every mineral-year.",
        "pass_rule": "Mass-balance flags must all pass and supply shares must be bounded and sum to 1.",
    },
    {
        "layer": "Numerical integrity",
        "purpose": "Detect model-breakage symptoms before interpreting results.",
        "evidence": "Scenario outputs are screened for NaN, infinity and negative values in key economic and environmental series.",
        "pass_rule": "No invalid values and no negative jobs, GVA, output or CO2.",
    },
    {
        "layer": "Repeatability",
        "purpose": "Make sure results are reproducible for the same configuration and seed.",
        "evidence": "Two identical runs are compared on cumulative GVA and the full numeric time series.",
        "pass_rule": "Cumulative GVA and time-series arrays must match within numerical tolerance.",
    },
    {
        "layer": "Economic structure",
        "purpose": "Check the SAM, CGE and multiplier logic against benchmark expectations.",
        "evidence": "SAM row/column balance, mining GVA anchor, CGE no-shock wage replication and multiplier ranges.",
        "pass_rule": "Balance error near zero, benchmark replication successful and multipliers in plausible ranges.",
    },
    {
        "layer": "Behavioural features",
        "purpose": "Verify that policy and shock mechanisms move in the intended direction.",
        "evidence": "Stockpile depletion, diversification exposure, time-varying shock onset and impact-layer response.",
        "pass_rule": "Directional tests must pass without violating reserves, shares or output bounds.",
    },
    {
        "layer": "Stress testing",
        "purpose": "Exercise the model across random but valid policy bundles and shocks.",
        "evidence": "Thirty fuzz configurations vary policies, demand growth, import caps, price paths and CGE use.",
        "pass_rule": "Every random configuration must run and preserve the core invariants.",
    },
]


def _run_validation():
    verify_model.RESULTS.clear()
    log = io.StringIO()
    with contextlib.redirect_stdout(log):
        print("MODEL VERIFICATION & VALIDATION")
        for _, test_func in TEST_GROUPS:
            try:
                test_func()
            except Exception as exc:
                verify_model.check(
                    f"{test_func.__name__} ran without error",
                    False,
                    f"EXCEPTION: {exc!r}",
                )
    rows = []
    for passed, name, detail in verify_model.RESULTS:
        rows.append({
            "status": "PASS" if passed else "FAIL",
            "check": name,
            "detail": detail,
        })
    return rows, log.getvalue()


def _write_results(rows, log_text):
    os.makedirs(OUT, exist_ok=True)
    df = pd.DataFrame(rows)
    csv_path = os.path.join(OUT, "model_validation_results.csv")
    json_path = os.path.join(OUT, "model_validation_results.json")
    log_path = os.path.join(OUT, "model_validation_console.log")
    df.to_csv(csv_path, index=False)
    payload = {
        "run_timestamp": datetime.now().isoformat(timespec="seconds"),
        "checks_total": len(rows),
        "checks_passed": int((df["status"] == "PASS").sum()),
        "checks_failed": int((df["status"] == "FAIL").sum()),
        "checks": rows,
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2)
    with open(log_path, "w", encoding="utf-8") as f:
        f.write(log_text)
    return csv_path, json_path, log_path, payload


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def _add_title(doc, payload):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("NI Circular Minerals Model Validation Report")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    r = meta.add_run(
        f"Generated {payload['run_timestamp']} | "
        f"{payload['checks_passed']} passed, {payload['checks_failed']} failed, "
        f"{payload['checks_total']} checks total"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")


def _add_summary_table(doc, payload):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Result", "Passed", "Failed", "Total checks"]
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True)
        _set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    row = table.add_row().cells
    result = "PASS" if payload["checks_failed"] == 0 else "FAIL"
    values = [result, payload["checks_passed"], payload["checks_failed"], payload["checks_total"]]
    for idx, value in enumerate(values):
        _set_cell_text(row[idx], value, bold=(idx == 0), color=("1F7A1F" if result == "PASS" else "9B1C1C"))


def _add_design(doc):
    doc.add_heading("How the validation is designed", level=1)
    doc.add_paragraph(
        "The validation is layered so that the model is tested from the inside out: "
        "external calibration first, then physical and accounting identities, then "
        "numerical integrity, reproducibility, economic benchmark replication, policy "
        "behaviour and randomized stress testing. This combination is intended to catch "
        "both calibration drift and hidden implementation failures."
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1.35, 1.75, 2.25, 1.15]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    for idx, header in enumerate(["Layer", "Purpose", "Evidence used", "Pass rule"]):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True)
        _set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for item in VALIDATION_DESIGN:
        cells = table.add_row().cells
        _set_cell_text(cells[0], item["layer"], bold=True)
        _set_cell_text(cells[1], item["purpose"])
        _set_cell_text(cells[2], item["evidence"])
        _set_cell_text(cells[3], item["pass_rule"])


def _add_failures(doc, rows):
    failures = [r for r in rows if r["status"] == "FAIL"]
    doc.add_heading("Validation result", level=1)
    if not failures:
        doc.add_paragraph(
            "All implemented checks passed in this run. The result supports using the "
            "current model outputs for scenario-comparison and policy-design analysis, "
            "subject to the data limitations documented in the model."
        )
        return

    doc.add_paragraph(
        "One or more checks failed. Treat affected model outputs as provisional until "
        "the failures below are fixed and this validation is rerun."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for idx, header in enumerate(["Failed check", "Detail"]):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True)
        _set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for item in failures:
        cells = table.add_row().cells
        _set_cell_text(cells[0], item["check"], bold=True, color="9B1C1C")
        _set_cell_text(cells[1], item["detail"])


def _add_results_by_group(doc, rows):
    doc.add_heading("Check-level results", level=1)
    doc.add_paragraph(
        "The table below lists every automated check, its status and the diagnostic "
        "detail produced by the validation harness."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate(["Status", "Check", "Diagnostic detail"]):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True)
        _set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for item in rows:
        cells = table.add_row().cells
        color = "1F7A1F" if item["status"] == "PASS" else "9B1C1C"
        _set_cell_text(cells[0], item["status"], bold=True, color=color)
        _set_cell_text(cells[1], item["check"])
        _set_cell_text(cells[2], item["detail"] or "")


def _add_limitations(doc):
    doc.add_heading("Interpretation and limitations", level=1)
    items = [
        "Passing validation means the implemented equations, accounting identities and benchmark checks are internally consistent; it does not turn proxy inputs into forecasts.",
        "The strongest external validation is currently the Minviro anchor comparison for the I-O core. Several CGE, SAM, company-score and local-flow parameters remain proxy or desk-estimated inputs.",
        "The fuzz tests improve confidence that valid policy bundles do not break invariants, but they do not prove every possible real-world configuration has been calibrated.",
        "The recommended validation cycle is to rerun this script after any change to model equations, parameters, company data, scenario definitions or output schemas.",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def _write_docx(rows, payload):
    os.makedirs(WORD, exist_ok=True)
    doc = Document()
    _style_doc(doc)
    _add_title(doc, payload)
    _add_summary_table(doc, payload)
    _add_design(doc)
    _add_failures(doc, rows)
    _add_results_by_group(doc, rows)
    _add_limitations(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("NI Circular Minerals Model Validation")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string("555555")

    path = os.path.join(WORD, "Model_Validation_Report.docx")
    doc.save(path)
    return path


def main():
    rows, log_text = _run_validation()
    csv_path, json_path, log_path, payload = _write_results(rows, log_text)
    docx_path = _write_docx(rows, payload)
    print(f"Validation checks: {payload['checks_passed']} passed, {payload['checks_failed']} failed, {payload['checks_total']} total")
    print(f"Wrote {os.path.relpath(csv_path, ROOT)}")
    print(f"Wrote {os.path.relpath(json_path, ROOT)}")
    print(f"Wrote {os.path.relpath(log_path, ROOT)}")
    print(f"Wrote {os.path.relpath(docx_path, ROOT)}")
    return 1 if payload["checks_failed"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
